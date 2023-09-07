import re
import os, glob
import comtypes.client
import docx
from nltk.stem import WordNetLemmatizer
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import win32com.client as win32


# Data 
country_data = {'Afghanistan': 'Afghan',
 'Albania': 'Albanian',
 'Algeria': 'Algerian',
 'Argentina': 'Argentine',
 'Australia': 'Australian',
 'Austria': 'Austrian',
 'Bangladesh': 'Bangladeshi',
 'Belgium': 'Belgian',
 'Bolivia': 'Bolivian',
 'Botswana': 'Batswana',
 'Brazil': 'Brazilian',
 'Bulgaria': 'Bulgarian',
 'Cambodia': 'Cambodian',
 'Cameroon': 'Cameroonian',
 'Canada': 'Canadian',
 'Chile': 'Chilean',
 'China': 'Chinese',
 'Colombia *': 'Colombian',
 'Costa Rica': 'Costa Rican',
 'Croatia': 'Croatian',
 'Cuba': 'Cuban',
 'Czech Republic': 'Czech',
 'Denmark': 'Danish',
 'Dominican Republic': 'Dominican',
 'Ecuador': 'Ecuadorian',
 'Egypt': 'Egyptian',
 'El Salvador': 'Salvadorian',
 'England': 'English',
 'Estonia': 'Estonian',
 'Ethiopia': 'Ethiopian',
 'Fiji': 'Fijian',
 'Finland': 'Finnish',
 'France': 'French',
 'Germany': 'German',
 'Ghana': 'Ghanaian',
 'Greece': 'Greek',
 'Guatemala': 'Guatemalan',
 'Haiti': 'Haitian',
 'Honduras': 'Honduran',
 'Hungary': 'Hungarian',
 'Iceland': 'Icelandic',
 'India': 'Indian',
 'Indonesia': 'Indonesian',
 'Iran': 'Iranian',
 'Iraq': 'Iraqi',
 'Ireland': 'Irish',
 'Israel': 'Israeli',
 'Italy': 'Italian',
 'Jamaica': 'Jamaican',
 'Japan': 'Japanese',
 'Jordan': 'Jordanian',
 'Kenya': 'Kenyan',
 'Kuwait': 'Kuwaiti',
 'Laos': 'Lao',
 'Latvia': 'Latvian',
 'Lebanon': 'Lebanese',
 'Libya': 'Libyan',
 'Lithuania': 'Lithuanian',
 'Madagascar': 'Malagasy',
 'Malaysia': 'Malaysian',
 'Mali': 'Malian',
 'Malta': 'Maltese',
 'Mexico': 'Mexican',
 'Mongolia': 'Mongolian',
 'Morocco': 'Moroccan',
 'Mozambique': 'Mozambican',
 'Namibia': 'Namibian',
 'Nepal': 'Nepalese',
 'Netherlands': 'Dutch',
 'New Zealand': 'New Zealand',
 'Nicaragua': 'Nicaraguan',
 'Nigeria': 'Nigerian',
 'Norway': 'Norwegian',
 'Pakistan': 'Pakistani',
 'Panama': 'Panamanian',
 'Paraguay': 'Paraguayan',
 'Peru': 'Peruvian',
 'Philippines': 'Philippine',
 'Poland': 'Polish',
 'Portugal': 'Portuguese',
 'Romania': 'Romanian',
 'Russia': 'Russian',
 'Saudi Arabia': 'Saudi',
 'Scotland': 'Scottish',
 'Senegal': 'Senegalese',
 'Serbia': 'Serbian',
 'Singapore': 'Singaporean',
 'Slovakia': 'Slovak',
 'South Africa': 'South African',
 'South Korea': 'Korean',
 'Spain': 'Spanish',
 'Sri Lanka': 'Sri Lankan',
 'Sudan': 'Sudanese',
 'Sweden': 'Swedish',
 'Switzerland': 'Swiss',
 'Syria': 'Syrian',
 'Taiwan': 'Taiwanese',
 'Tajikistan': 'Tajikistani',
 'Thailand': 'Thai',
 'Tonga': 'Tongan',
 'Tunisia': 'Tunisian',
 'Turkey': 'Turkish',
 'Ukraine': 'Ukrainian',
 'United Arab Emirates': 'Emirati',
 '(The) United Kingdom': 'British',
 '(The) United States': 'American **',
 'Uruguay': 'Uruguayan',
 'Venezuela': 'Venezuelan',
 'Vietnam': 'Vietnamese',
 'Wales': 'Welsh',
 'Zambia': 'Zambian',
 'Zimbabwe': 'Zimbabwean',
 'UAE': 'Emirati',
 'UK': 'British',
 'USA': 'American',
 }

#regex data

date_pattern = r'\d{1,2}(st|nd|rd|th|TH|ST|ND|RD)?[\s/]*\w+[\s/]\d{4}|\d{1,2}-\d{1,2}-\d{4}'



formats = ["%b %d, %Y", "%Y-%m-%d", "%m/%d/%Y"]

phone_pattern = re.compile(r"(\d{3}[-.\s]??\d{3}[-.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-.\s]??\d{4})|"r"(\+\d{2}\s\d{4}\s\d{4})|"r"(\+\d{1,12}(\s\d{3}){0,3})")

passport_pattern = re.compile(r"\b[A-Z]+\d+[A-Z\d]*\b")

name_pattern = re.compile(r"\b[a-zA-Z][A-Za-z]+ [a-zA-Z][A-Za-z]+\b")


def file_conv(dir_path):
    for file in glob.glob(os.path.join(dir_path, "*.rtf")):
        root, ext = os.path.splitext(file)
        if ext == ".rtf":
            new_path = root + ".doc"
            os.rename(file, new_path)
    files = os.listdir(dir_path)
    non_inluded = []
    for f in files:
        print(f"{f} - Conversion Initiated")
        try:
            if f.endswith('.docx'):
                pass
            elif f.endswith('.doc'):
                input_file = os.path.join(dir_path,f)
                output_file = os.path.join(dir_path, f.replace('.doc', '.docx'))
                word = comtypes.client.CreateObject("Word.Application")
                doc = word.Documents.Open(os.path.abspath(input_file))
                doc.SaveAs(output_file, FileFormat=12)
                doc.Close()
                word.Quit()
                print(f , ": doc conversion done")

            elif f.endswith('.rtf') :
                word = win32.Dispatch("Word.Application")
                doc = word.Documents.Open(os.path.join(dir_path,f))
                doc.SaveAs(os.path.join(dir_path, f.replace('.rtf', '.docx')), FileFormat=12)
                doc.Close()
                word.Quit()
                print(f , ": rtf conversion done")
        except:
            print(f , ": this rtf conversion failed reasons may be - File format not supported / read only document format")
            non_inluded.append(f)
    print(f'Files not included in modification : {non_inluded}')

class Edit_doc():
    def __init__(self, new_in_path, out_path, name_pattern, new_name, emp_id, new_email, address):
        self.in_path = new_in_path
        self.out_path = out_path
        self.doc = docx.Document(new_in_path)
        self.name_pattern = name_pattern
        self.new_name = new_name
        self.emp_id = emp_id
        self.new_email = new_email
        self.address = address

    def edit_doc(self):
        name_list = []
        for para in self.doc.paragraphs[:1]:
            for run in para.runs:
                names = self.name_pattern.findall(run.text)
                for name in names:
                    run.text = run.text.replace(name, self.new_name)
                    name_list.append(name)
        for para in self.doc.paragraphs:
            for run in para.runs:
                for n in name_list:
                    for n in n.split(" "):
                        if n.lower() in run.text.split(" ") or n.capitalize() in run.text.split(" ") or n.upper() in run.text.split(" ") :
                            run.text = para.text.replace(n.lower(), self.new_name.split(" ")[1].capitalize())
                            run.text = para.text.replace(n.upper(), self.new_name.split(" ")[1].capitalize())
                            run.text = para.text.replace(n.capitalize(), self.new_name.split(" ")[1].capitalize())
        
        header = self.doc.sections[0].header
        paragraph = header.paragraphs
        for para in paragraph:
            for f_n in name_list:
                for n in f_n.split(" "):
                    if f_n in para.text:
                        para.text = para.text.replace(f_n, self.new_name)
                    if n in para.text:
                        para.text = para.text.replace(n, self.new_name.split(" ")[1])

        country_names = country_data.keys()

        for country in country_names:
            for para in self.doc.paragraphs[:5]+self.doc.paragraphs[-20:]:
                if country in [i.capitalize().strip() for i in para.text.split(" ")]:
                    para.text = self.address
                lemetizr = WordNetLemmatizer()
                if country_data[country] in [ lemetizr.lemmatize(i) for i in para.text.split(" ")]:
                    para.text = para.text.replace(country_data[country].capitalize(), "UK")

        for para in self.doc.paragraphs[:5]+self.doc.paragraphs[-20:]:
            para.text = re.sub(phone_pattern, self.emp_id , para.text)

            passports = passport_pattern.findall(para.text)
            for p in passports:
                if len(p) > 6 and len(p) < 12 :
                    para.text = para.text.replace(p, '#'*6)

            para.text = re.sub(date_pattern, '##/##/####', para.text)

            rels = self.doc.part.rels
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

            for rel in rels:
                if rels[rel].reltype == RT.HYPERLINK:
                    rels[rel]._target =''
            if re.match(email_pattern, para.text):
                self.doc.add_paragraph (self.new_email)

        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.text or run.element.xml.startswith('<w:drawing'):
                    continue
                else:
                    run.element.clear()

        self.doc.save(self.out_path)
        return self.out_path


def csv_json(csv_file,json_file = "output.json"):
    import csv
    import json
    csv_file = csv_file
    json_file = json_file
    with open(csv_file, mode='r') as csv_file:
        csv_reader = csv.DictReader(csv_file)
        data_list = []
        for row in csv_reader:
            data_list.append(row)
    with open(json_file, mode='w') as json_file:
        json.dump(data_list, json_file, indent=4)
    print('Done')